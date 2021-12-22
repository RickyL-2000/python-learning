# %%
import os
from docx import Document
import jieba

DISPLAY_LIMIT = 40

def _check_suffix(path, suffix):
    filename = path.split('/')[-1]
    if len(filename) == 0 or len(filename) < len(suffix):
        return False
    else:
        if len(suffix) == 0:
            return True
        elif filename[-len(suffix):] != suffix:
            return False
        else:
            return True

def find_files(file_list: list, path="./", suffix=""):
    if os.path.isfile(path):
        if _check_suffix(path, suffix):
            file_list.append(path)
    else:
        for f_path in os.listdir(path):
            find_files(file_list, os.path.join(path, f_path), suffix)

def parse_index(file_list, msg):
    """对输入的序号列表进行分析

    Args:
        file_list (list): 该目录下待选择的文件列表
        msg (str): 待分析的输入

    Returns:
        int: 状态序号 (-2代表运行错误，-1代表检查错误，0代表运行成功)
        list: 被选择的文件序号列表
    """
    try:
        msgs = msg.split(" ")
        if msgs[0] != "f":
            return -1, None
        else:
            idx_lst = []
            for i in range(1, len(msgs)):
                if "-" in msgs[i]:
                    begin, end = msgs[i].split("-")
                    if begin > end:
                        return -1, None
                    idx_lst.extend(list(range(int(begin), int(end)+1)))
                else:
                    idx_lst.append(int(msgs[i]))
            idx_lst = list(set(idx_lst))
            if max(idx_lst) >= len(file_list):
                return -1, None
            return 0, idx_lst
    except:
        return -2, None

def process_input(file_list):
    print("该目录下所有符合要求的文件如下:")
    page_num = 0
    page_size = DISPLAY_LIMIT
    print("-"*40)
    print(f"页码: {page_num}")
    print(f"[显示序号 0-{min(len(file_list), DISPLAY_LIMIT)}]")
    for i in range(page_num*page_size, min((page_num+1)*page_size, len(file_list))):
        print(f"{i}".ljust(5) + file_list[i])

    target_idxs = []
    
    while True:
        print("-"*40)
        print("请选择您需要操作的文件:")
        msg = input("输入: ")
        if msg == "q":
            return 0, None
        elif msg == "*":
            return 0, list(range(len(file_list)))
        elif msg == "ok":
            return 0, target_idxs
        elif msg == "u":
            print("-"*40)
            if page_num == 0: print("已经是第一页！")
            page_num = max(0, page_num-1)
            print(f"页码: {page_num}")
            print(f"[显示序号 {page_num*page_size}-{min((page_num+1)*page_size, len(file_list))}]")
            for i in range(page_num*page_size, min((page_num+1)*page_size, len(file_list))):
                print(f"{i}".ljust(5) + file_list[i])
            continue
        elif msg == "d":
            print("-"*40)
            if page_num == 0: print("已经是最后一页！")
            page_num = min(len(file_list) // DISPLAY_LIMIT, page_num+1)
            print(f"页码: {page_num}")
            print(f"[显示序号 {page_num*page_size}-{min((page_num+1)*page_size, len(file_list))}]")
            for i in range(page_num*page_size, min((page_num+1)*page_size, len(file_list))):
                print(f"{i}".ljust(5) + file_list[i])
            continue
        else:
            _status, idx_lst = parse_index(file_list, msg)
            if _status == 0:
                target_idxs.extend(idx_lst)
                target_idxs = list(set(target_idxs))
            else:
                print("无效输入")
        
def program_exit():
    print("-"*40)
    print("程序完成")
    exit()

def combine_py(file_list):
    document = Document()

    for filename in file_list:
        with open(filename, "r", encoding="utf-8") as f:
            content = f.read()
        document.add_heading(f'file path: {filename}', level=1)
        document.add_paragraph(content)
    
    if not os.path.exists("output"):
        os.makedirs("output")
    output_path = os.path.join("output", "all_python.docx")
    document.save(output_path)
    print(f"已写入文件 {output_path}")

def wordfreq(file_list):
    word_list = []
    word_dict = {}

    for filename in file_list:
        with open(filename, "r", encoding="utf-8") as f:
            content = f.read()
        word_list.extend(jieba.lcut(content))

    for w in word_list:
        if len(w) == 1:
            continue
        if w in word_dict:
            word_dict[w] += 1
        else:
            word_dict[w] = 1
    
    items_list = sorted(list(word_dict.items()), key=lambda x: x[1], reverse=True)
    
    print("展示前10个最高频的词语的频率:")
    for i in range(10):
        print(items_list[i][0], items_list[i][1])

    if not os.path.exists("output"):
        os.makedirs("output")
    output_path = os.path.join("output", "词频.csv")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('word, freq\n')
        for i in range(len(items_list)):
            f.write(items_list[i][0] + ',' + str(items_list[i][1]) + '\n')
    
    print(f"已写入文件 {output_path}")
    

def main():
    print("程序开始运行...")
    print("-"*40+"\n说明:")
    print(f"本程序默认展示前{DISPLAY_LIMIT}个文件。")
    print(f"若文件数量超过{DISPLAY_LIMIT}，则需要输入u(page up)或者d(page down)来进行翻页")
    print("您可以输入'*'来表示选择全部文件，或者指定特定文件进行处理")
    print("指定特定文件，输入由字符 'f' 开头，后接文件序号。连续文件序号可以用连字符'-'连接(前后都包括，且不能后小于前)。每个序号(组)之间用空格连接。")
    print("示例如下:")
    print("    f 1 2-3 6 18-30")
    print("该语句可以多次输入，直到输入'ok'表示输入完毕，程序继续运行；或者输入'*'表示直接全选；或者'q'表示退出，程序完成")
    print("-"*40)
    print("请选择您想要的功能(输入1或者2，输入q则退出程序)")
    print("1. 合并python文件")
    print("2. 统计中文文本文件的词频")
    print("-"*40)
    while True:
        msg = input("输入: ")
        if msg == "q":
            program_exit()
        elif msg == "1":
            print("-"*40)
            print("您选择合并python文件")

            file_list = []
            find_files(file_list, path="data", suffix="py")

            _status, target_idxs = process_input(file_list)
            if _status == 0 and target_idxs is None:
                program_exit()
            
            target_list = [file_list[idx] for idx in target_idxs]

            # print("Your chosen file indices:")
            # print(target_idxs)

            combine_py(target_list)
            program_exit()

        elif msg == "2":
            print("-"*40)
            print("您选择中文词频统计")

            file_list = []
            find_files(file_list, path="data", suffix="txt")

            _status, target_idxs = process_input(file_list)
            if _status == 0 and target_idxs is None:
                program_exit()
            
            target_list = [file_list[idx] for idx in target_idxs]

            # print(target_idxs)

            wordfreq(target_list)
            program_exit()

        else:
            print("无效输入")
            print("-"*40)

if __name__ == '__main__':
    main()
